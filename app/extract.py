"""Document ingestion.

Turns an uploaded file into a uniform list of pages. Every downstream stage
sees the same shape regardless of source format, so the chunker never needs to
know whether the text came from a PDF, a Word document, or a photograph.

Supported: PDF, DOCX, TXT/MD, and images (PNG, JPEG, WEBP, GIF, BMP, TIFF).

Images are read by the multimodal model at ingest rather than by a separate OCR
engine. That keeps the dependency surface small and handles cases classical OCR
does poorly -- diagrams, charts, screenshots, handwriting -- because the model
describes structure as well as transcribing glyphs. The same path rescues
scanned PDFs, which carry no text layer for pypdf to read.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

import docx
from pypdf import PdfReader

from app.config import (
    IMAGE_EXTENSIONS,
    IMAGE_MIME_TYPES,
    PDF_MIN_CHARS_PER_PAGE,
    SUPPORTED_EXTENSIONS,
)

# Asks for a faithful transcription plus enough description that a chart or
# diagram becomes searchable text. Explicitly forbids commentary, because
# anything the model adds here becomes indexed "source" material that a later
# answer could cite as though it came from the document.
IMAGE_INSTRUCTION = """\
Transcribe ALL text visible in this image, exactly as written, preserving \
the reading order and any table structure.

If the image contains a chart, diagram, form, or screenshot, also describe \
its content and any data it shows, so the information is searchable.

Output only the transcription and description. Do not add commentary, \
interpretation, or notes about the image quality.\
"""

PDF_INSTRUCTION = """\
Transcribe ALL text in this document, exactly as written, preserving reading \
order and table structure. Mark each page boundary on its own line as:

[[PAGE n]]

Output only the transcription. Do not add commentary or interpretation.\
"""

# The model emits this marker so page provenance survives vision extraction.
_PAGE_MARKER = "[[PAGE"


@dataclass
class Page:
    """One unit of source text with the provenance needed for citation.

    `page_number` is 1-indexed. Formats without real pagination (DOCX, TXT,
    a single image) report page 1 -- the field always carries something citable.
    """

    text: str
    page_number: int


class UnsupportedFileError(ValueError):
    """Raised for a file type this build does not handle."""


class ExtractionError(RuntimeError):
    """Raised when a supported file could not be read."""


def extract(filename: str, data: bytes) -> list[Page]:
    """Extract text from an uploaded file.

    Args:
        filename: Original filename -- its extension selects the parser.
        data: Raw file bytes.

    Returns:
        Pages with non-empty text, in document order.

    Raises:
        UnsupportedFileError: Extension is not supported.
        ExtractionError: A supported file could not be read.
    """
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileError(
            f"Cannot process '{filename}'. Supported types: {supported}"
        )

    if suffix in IMAGE_EXTENSIONS:
        pages = _extract_image(data, IMAGE_MIME_TYPES[suffix])
    elif suffix == ".pdf":
        pages = _extract_pdf(data)
    elif suffix == ".docx":
        pages = _extract_docx(data)
    else:
        pages = _extract_text(data)

    return [p for p in pages if p.text.strip()]


def _extract_image(data: bytes, mime_type: str) -> list[Page]:
    """Read an image with the vision model."""
    from app.llm import LLMError, get_llm

    try:
        text = get_llm().extract_from_media(data, mime_type, IMAGE_INSTRUCTION)
    except LLMError as exc:
        raise ExtractionError(
            f"Could not read the image: {exc}. Image ingestion needs Vertex AI "
            "to be reachable."
        ) from exc

    if not text.strip():
        raise ExtractionError(
            "No text or describable content was found in this image."
        )

    return [Page(text=text, page_number=1)]


def _extract_pdf(data: bytes) -> list[Page]:
    """Extract per page, falling back to vision for scanned documents.

    The text layer is tried first: it is free, instant, and exact. A scan has
    no text layer, so a PDF that yields almost nothing is re-read with the
    vision model rather than indexed as a near-empty document.
    """
    reader = PdfReader(io.BytesIO(data))
    pages: list[Page] = []

    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            # One malformed page shouldn't lose the rest of the document.
            text = ""
        pages.append(Page(text=text, page_number=index))

    extracted = sum(len(p.text.strip()) for p in pages)
    threshold = PDF_MIN_CHARS_PER_PAGE * max(len(pages), 1)

    if extracted >= threshold:
        return pages

    # Too little text for the page count -- treat as scanned.
    try:
        return _extract_pdf_via_vision(data)
    except ExtractionError:
        # Vision unavailable: return whatever the text layer gave us. If that
        # is nothing, the caller reports an empty document, which is accurate.
        return pages


def _extract_pdf_via_vision(data: bytes) -> list[Page]:
    """Read a scanned PDF with the vision model, preserving page numbers."""
    from app.llm import LLMError, get_llm

    try:
        text = get_llm().extract_from_media(data, "application/pdf", PDF_INSTRUCTION)
    except LLMError as exc:
        raise ExtractionError(f"Could not read the scanned PDF: {exc}") from exc

    if not text.strip():
        raise ExtractionError("No text was found in this PDF.")

    return _split_page_markers(text)


def _split_page_markers(text: str) -> list[Page]:
    """Split vision output on [[PAGE n]] markers into numbered pages.

    Falls back to a single page if the model omitted the markers -- losing page
    granularity is acceptable, losing the content is not.
    """
    if _PAGE_MARKER not in text:
        return [Page(text=text, page_number=1)]

    pages: list[Page] = []
    current_number = 1
    buffer: list[str] = []

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(_PAGE_MARKER) and stripped.endswith("]]"):
            if buffer:
                pages.append(
                    Page(text="\n".join(buffer).strip(), page_number=current_number)
                )
                buffer = []
            digits = "".join(ch for ch in stripped if ch.isdigit())
            current_number = int(digits) if digits else current_number + 1
        else:
            buffer.append(line)

    if buffer:
        pages.append(Page(text="\n".join(buffer).strip(), page_number=current_number))

    return pages


def _extract_docx(data: bytes) -> list[Page]:
    """Extract paragraphs and table cells.

    Table content is included deliberately: a paragraph-only read silently
    drops tables, which in business documents often hold exactly the facts
    someone will ask about -- pricing, dates, specifications.
    """
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                # Pipe-joined so the row reads as one related record rather
                # than dissolving into unrelated fragments.
                parts.append(" | ".join(cells))

    return [Page(text="\n\n".join(parts), page_number=1)]


def _extract_text(data: bytes) -> list[Page]:
    """Decode plain text, tolerating non-UTF-8 files rather than failing."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")

    return [Page(text=text, page_number=1)]
