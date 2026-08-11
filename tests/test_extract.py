"""Extraction tests — format handling, image ingestion, and type boundaries."""

import io
from unittest.mock import MagicMock, patch

import docx
import pytest

from app.extract import ExtractionError, UnsupportedFileError, extract


def _mock_pdf(page_texts: list[str]):
    """Patch pypdf to yield the given per-page text.

    Authoring a real PDF with a text layer needs a rendering library; mocking
    the reader targets the branch under test — whether the text layer is
    considered sufficient — without that dependency.
    """
    pages = []
    for text in page_texts:
        page = MagicMock()
        page.extract_text.return_value = text
        pages.append(page)

    reader = MagicMock()
    reader.pages = pages
    return patch("app.extract.PdfReader", return_value=reader)


def test_txt_extraction():
    pages = extract("notes.txt", b"Hello world.\n\nSecond paragraph.")

    assert len(pages) == 1
    assert "Hello world." in pages[0].text
    assert pages[0].page_number == 1


def test_non_utf8_text_does_not_crash():
    """A latin-1 file should degrade gracefully rather than fail the upload."""
    pages = extract("legacy.txt", "Café résumé".encode("latin-1"))

    assert len(pages) == 1
    assert pages[0].text


def test_docx_extracts_paragraphs():
    document = docx.Document()
    document.add_paragraph("The Enterprise plan costs $99/month.")
    document.add_paragraph("Support is included.")

    buffer = io.BytesIO()
    document.save(buffer)

    pages = extract("plans.docx", buffer.getvalue())

    assert len(pages) == 1
    assert "$99/month" in pages[0].text
    assert "Support is included." in pages[0].text


def test_docx_includes_table_content():
    """Tables often hold exactly the facts users ask about."""
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Plan"
    table.cell(0, 1).text = "Price"
    table.cell(1, 0).text = "Enterprise"
    table.cell(1, 1).text = "$99"

    buffer = io.BytesIO()
    document.save(buffer)

    pages = extract("pricing.docx", buffer.getvalue())

    assert "Enterprise" in pages[0].text
    assert "$99" in pages[0].text


def test_unsupported_extension_is_rejected():
    with pytest.raises(UnsupportedFileError) as exc:
        extract("archive.zip", b"PK\x03\x04")

    # The error should name what IS supported, not just what failed.
    assert ".pdf" in str(exc.value)


def test_empty_pages_are_dropped():
    assert extract("blank.txt", b"   \n\n   ") == []


# --- Image ingestion ----------------------------------------------------------


def _mock_vision(text: str):
    """Patch the vision model to return a fixed transcription."""
    llm = MagicMock()
    llm.extract_from_media.return_value = text
    return patch("app.llm.get_llm", return_value=llm), llm


def test_image_is_transcribed_by_vision_model():
    patcher, llm = _mock_vision("INVOICE\nTotal: $4,200\nDue: 3 March 2026")

    with patcher:
        pages = extract("invoice.png", b"\x89PNG\r\n\x1a\n")

    assert len(pages) == 1
    assert "$4,200" in pages[0].text
    assert pages[0].page_number == 1

    # The image bytes and the correct MIME type must reach the model.
    kwargs = llm.extract_from_media.call_args
    assert kwargs.args[1] == "image/png"


@pytest.mark.parametrize(
    "filename,mime",
    [
        ("photo.jpg", "image/jpeg"),
        ("photo.jpeg", "image/jpeg"),
        ("shot.webp", "image/webp"),
        ("scan.tiff", "image/tiff"),
        ("chart.gif", "image/gif"),
        ("art.bmp", "image/bmp"),
    ],
)
def test_image_formats_map_to_correct_mime_type(filename, mime):
    patcher, llm = _mock_vision("some text")

    with patcher:
        extract(filename, b"binary-image-data")

    assert llm.extract_from_media.call_args.args[1] == mime


def test_image_with_no_content_raises():
    patcher, _ = _mock_vision("   ")

    with patcher:
        with pytest.raises(ExtractionError) as exc:
            extract("blank.png", b"binary")

    assert "no text" in str(exc.value).lower()


def test_image_extraction_failure_is_actionable():
    """A vision failure should explain itself, not surface a raw traceback."""
    from app.llm import LLMError

    llm = MagicMock()
    llm.extract_from_media.side_effect = LLMError("credentials missing")

    with patch("app.llm.get_llm", return_value=llm):
        with pytest.raises(ExtractionError) as exc:
            extract("photo.png", b"binary")

    assert "Vertex AI" in str(exc.value)


# --- Scanned PDF fallback -----------------------------------------------------


def test_pdf_with_text_layer_skips_vision():
    """The text layer is free and exact — vision must not be called for it."""
    llm = MagicMock()

    with _mock_pdf(["Quarterly revenue was 4.2 million dollars. " * 5]):
        with patch("app.llm.get_llm", return_value=llm):
            pages = extract("report.pdf", b"%PDF-fake")

    assert any("4.2 million" in p.text for p in pages)
    llm.extract_from_media.assert_not_called()


def test_scanned_pdf_falls_back_to_vision():
    """A PDF with no text layer would otherwise index as empty."""
    patcher, llm = _mock_vision(
        "[[PAGE 1]]\nScanned contract text\n[[PAGE 2]]\nSignatures"
    )

    with _mock_pdf(["", ""]):  # two pages, no text layer
        with patcher:
            pages = extract("scan.pdf", b"%PDF-fake")

    assert llm.extract_from_media.called
    assert llm.extract_from_media.call_args.args[1] == "application/pdf"
    assert len(pages) == 2
    assert pages[0].page_number == 1
    assert "Scanned contract" in pages[0].text
    assert pages[1].page_number == 2


def test_pdf_with_trivial_text_is_treated_as_scanned():
    """A page yielding only a stray page number is a scan, not a document."""
    patcher, llm = _mock_vision("Full transcription of the scan")

    with _mock_pdf(["3", "4"]):
        with patcher:
            pages = extract("scan.pdf", b"%PDF-fake")

    assert llm.extract_from_media.called
    assert "Full transcription" in pages[0].text


def test_vision_output_without_page_markers_still_works():
    """Losing page granularity is acceptable; losing content is not."""
    patcher, _ = _mock_vision("Text with no page markers at all")

    with _mock_pdf([""]):
        with patcher:
            pages = extract("scan.pdf", b"%PDF-fake")

    assert len(pages) == 1
    assert "no page markers" in pages[0].text


def test_scanned_pdf_without_vision_degrades_quietly():
    """If vision is unavailable, fall back to whatever the text layer gave."""
    from app.llm import LLMError

    llm = MagicMock()
    llm.extract_from_media.side_effect = LLMError("unavailable")

    with _mock_pdf([""]):
        with patch("app.llm.get_llm", return_value=llm):
            pages = extract("scan.pdf", b"%PDF-fake")

    # No exception; an empty result is reported accurately by the caller.
    assert pages == []
